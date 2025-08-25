#!/usr/bin/env python3
"""
Generate two-up DOCX cards for lamination.

Left QR: Support page (support.html)
Right QR: Room page (derived from QR image filename)

Logo source: images/nyu-logo.svg (converted to PNG at runtime for Word)
Output: qr_codes/print_cards.docx
"""
from pathlib import Path
import io
import re
import argparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[2]
QR_DIR = ROOT / "qr_codes"
IMG_DIR = ROOT / "images"
OUTPUT_DOCX = QR_DIR / "print_cards.docx"

# Pages base
BASE_URL = "https://mpaproomguide.github.io/MPAP-Room-Guide/"
SUPPORT_URL = BASE_URL + "support.html"

def read_png_bytes(png_path: Path) -> bytes:
    return png_path.read_bytes()

def list_room_qr_pngs():
    files = sorted([p for p in QR_DIR.glob("*.png") if "_QR" in p.stem])
    return files

def filename_to_room_url(p: Path) -> str:
    # Example: Room_307_QR.png -> room307.html (lowercase, remove spaces/underscores)
    name = p.stem.replace("_QR", "")
    # Heuristics: if it starts with "Room_", map to roomNNN.html; else, lower kebab
    if name.startswith("Room_"):
        slug = name.replace("Room_", "room").replace("_", "")
        return f"{slug}.html"
    else:
        slug = name.lower().replace(" ", "-").replace("_", "-")
        return f"{slug}.html"

def _zero_paragraph_spacing(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

def stylize_qr_image(qr_path: Path, bar_text: str = "SCAN ME") -> io.BytesIO:
    """Create an attractive QR image with rounded dark border and a bottom bar saying SCAN ME."""
    qr = Image.open(qr_path).convert("RGB")
    # Normalize QR size
    target_qr_side = 700
    qr = qr.resize((target_qr_side, target_qr_side), Image.LANCZOS)

    # Frame dimensions
    margin = 40
    bar_height = 120
    inner_radius = 70
    outer_radius = 75
    frame_w = target_qr_side + margin * 2
    frame_h = target_qr_side + margin * 2 + bar_height

    img = Image.new("RGB", (frame_w, frame_h), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Outer dark rounded rectangle
    draw.rounded_rectangle([(0, 0), (frame_w - 1, frame_h - 1)], radius=outer_radius, fill=(45, 45, 45))
    # Inner white rounded rectangle (content area, excluding bottom bar)
    inner_rect = [
        (margin // 2, margin // 2),
        (frame_w - margin // 2 - 1, frame_h - bar_height - margin // 2 - 1),
    ]
    draw.rounded_rectangle(inner_rect, radius=inner_radius, fill=(255, 255, 255))

    # Paste QR centered in inner area
    qr_x = (frame_w - target_qr_side) // 2
    qr_y = (inner_rect[0][1] + inner_rect[1][1] - target_qr_side) // 2
    img.paste(qr, (qr_x, qr_y))

    # Bottom bar with rounded bottom corners to match outer frame
    bar_top = frame_h - bar_height
    # Create a temporary image for the bar with rounded bottom corners
    bar_img = Image.new("RGBA", (frame_w, bar_height), (0, 0, 0, 0))
    bar_draw = ImageDraw.Draw(bar_img)
    bar_draw.rounded_rectangle([(0, -outer_radius), (frame_w - 1, bar_height - 1)], 
                                radius=outer_radius, fill=(45, 45, 45))
    
    # Text in bottom bar - draw on the bar image BEFORE pasting
    try:
        font_path = (ROOT / "fonts" / "NYUPerstare-VF.ttf")
        font = ImageFont.truetype(str(font_path), 85)
    except Exception:
        try:
            font = ImageFont.truetype("Arial Bold", 85)
        except:
            font = ImageFont.load_default()
    text_bbox = bar_draw.textbbox((0, 0), bar_text, font=font)
    text_w, text_h = text_bbox[2] - text_bbox[0], text_bbox[3] - text_bbox[1]
    text_x = (frame_w - text_w) // 2
    # Add padding from bottom - move text up more
    text_y = (bar_height - text_h) // 2 - 20
    # Use pure white color
    bar_draw.text((text_x, text_y), bar_text, fill="white", font=font)
    
    # Now paste the bar with text onto the main image
    img.paste(bar_img, (0, bar_top), bar_img)

    bio = io.BytesIO()
    img.save(bio, format="PNG")
    bio.seek(0)
    return bio

def add_card(document: Document, logo_png: bytes, support_qr: Path, room_qr: Path, room_label: str):
    table = document.add_table(rows=7, cols=2)
    table.autofit = False
    # Set column widths (~3.5in each for US Letter with margins)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(3.65)

    # Explicit row heights to ensure consistent order/appearance in Pages
    rows = table.rows
    rows[0].height_rule = WD_ROW_HEIGHT_RULE.AUTO  # Let Pages fit the content
    rows[1].height_rule = WD_ROW_HEIGHT_RULE.AUTO  # Let Pages fit the content
    rows[2].height = Inches(0.25)
    rows[2].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    rows[3].height = Inches(2.9)
    rows[3].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    rows[4].height = Inches(0.16)
    rows[4].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    rows[5].height = Inches(0.16)
    rows[5].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    rows[6].height = Inches(0.25)
    rows[6].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Row 0: Logo in right column only
    hdr_par = table.cell(0, 1).paragraphs[0]
    hdr_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_run = hdr_par.add_run()
    hdr_run.add_picture(io.BytesIO(logo_png), width=Inches(2.0))
    _zero_paragraph_spacing(hdr_par)

    # Row 1: Department name under logo
    dept_par = table.cell(1, 1).paragraphs[0]
    dept_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept_run = dept_par.add_run("MUSIC AND PERFORMING ARTS PROFESSIONS")
    dept_run.font.size = Pt(6)
    dept_run.font.name = 'NYU Perstare'
    dept_run.bold = False
    dept_run.font.weight = 500  # Medium weight
    _zero_paragraph_spacing(dept_par)

    # Row 2: Titles
    left_title = table.cell(2, 0).paragraphs[0]
    left_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_run = left_title.add_run("Support")
    left_run.bold = False
    left_run.font.name = 'NYU Perstare'
    _zero_paragraph_spacing(left_title)

    right_title = table.cell(2, 1).paragraphs[0]
    right_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_run = right_title.add_run("Room Instructions")
    right_run.bold = False
    right_run.font.name = 'NYU Perstare'
    _zero_paragraph_spacing(right_title)

    # Row 3: QRs
    lqr = table.cell(3, 0).paragraphs[0]
    lqr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lqr.add_run().add_picture(stylize_qr_image(support_qr), width=Inches(2.4))
    _zero_paragraph_spacing(lqr)
    rqr = table.cell(3, 1).paragraphs[0]
    rqr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rqr.add_run().add_picture(stylize_qr_image(room_qr), width=Inches(2.4))
    _zero_paragraph_spacing(rqr)

    # Row 4: Empty (no labels under QRs)
    lcap = table.cell(4, 0).paragraphs[0]
    _zero_paragraph_spacing(lcap)
    rcap = table.cell(4, 1).paragraphs[0]
    _zero_paragraph_spacing(rcap)

    # Row 5: Empty (spacing)
    lscan = table.cell(5, 0).paragraphs[0]
    _zero_paragraph_spacing(lscan)
    rscan = table.cell(5, 1).paragraphs[0]
    _zero_paragraph_spacing(rscan)

    # Row 6: Room name
    room_name = table.cell(6, 1).paragraphs[0]
    room_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    room_run = room_name.add_run(room_label)
    room_run.bold = True
    room_run.font.name = 'NYU Perstare'
    _zero_paragraph_spacing(room_name)
    return table

def main():
    parser = argparse.ArgumentParser(description="Generate two-up DOCX QR cards")
    parser.add_argument("--include", dest="includes", action="append", help="Label to include (repeat for multiple)")
    parser.add_argument("--output", dest="output", help="Output DOCX path (default: qr_codes/print_cards.docx)")
    args = parser.parse_args()
    include_labels = set(s.strip() for s in args.includes) if args.includes else None

    doc = Document()
    # Portrait US Letter, two stacked cards per page
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Apply NYU Perstare font if installed
    try:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'NYU Perstare'
        font.size = Pt(11)
        # Ensure name is honored in Word
        r = doc.add_paragraph().add_run()
        r.font.name = 'NYU Perstare'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'NYU Perstare')
        doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
    except Exception:
        pass

    # Ensure support QR exists; generate if missing
    support_qr = QR_DIR / "Support_QR.png"
    if not support_qr.exists():
        # Create a quick QR using the existing PNGs if any one exists as template; else fail gracefully
        try:
            import qrcode
            img = qrcode.make(SUPPORT_URL)
            img.save(support_qr)
        except Exception as e:
            raise SystemExit(f"Failed to generate Support QR: {e}")

    # Prefer user-provided logo, fallback to repo image
    custom_logo = Path("/Users/drewatz/Desktop/steinhardt cropped.png")
    logo_path = custom_logo if custom_logo.exists() else (IMG_DIR / "nyu-logo.png")
    logo_png_bytes = read_png_bytes(logo_path)

    room_pngs = list_room_qr_pngs()
    if not room_pngs:
        raise SystemExit("No room QR PNGs found in qr_codes/. Run generate_qr_codes.py first.")

    # Map filename to display label
    def label_from(p: Path) -> str:
        return p.stem.replace("_QR", "").replace("_", " ")

    if include_labels:
        room_pngs = [p for p in room_pngs if label_from(p) in include_labels]
    else:
        # Default set: Education, Studios, Film Cart
        education = {"Room 303","Room 304","Room 305","Room 306","Room 307","Room 770","Room 771","Room 777","Room 778","Room 779","Room 876","Room 985"}
        filmcart = {"Room 302 A","Room 302 B","Room 302 C","Room 302 D","6th Floor Conference Room","Room 774"}
        studios = {"Studio A","Studio C","Studio D","Studio D1","Studio E","Studio F"}
        allowed = education | filmcart | studios
        room_pngs = [p for p in room_pngs if label_from(p) in allowed]
    # Sort: Studios first A..F, then Education numeric, then Film cart
    def sort_key(p: Path):
        lbl = label_from(p)
        if lbl.startswith("Studio "):
            return (0, lbl)
        if lbl.startswith("Room "):
            try:
                num = int(''.join(ch for ch in lbl.split(' ')[1] if ch.isdigit()))
            except Exception:
                num = 9999
            return (1, num)
        if lbl.startswith("6th") or lbl == "Room 774":
            return (2, lbl)
        return (3, lbl)
    room_pngs.sort(key=sort_key)

    # Two different cards per page; editable cells throughout
    for idx in range(0, len(room_pngs), 2):
        p1 = room_pngs[idx]
        label1 = p1.stem.replace("_QR", "").replace("_", " ")
        add_card(doc, logo_png_bytes, support_qr, p1, label1)
        if idx + 1 < len(room_pngs):
            p2 = room_pngs[idx + 1]
            label2 = p2.stem.replace("_QR", "").replace("_", " ")
            add_card(doc, logo_png_bytes, support_qr, p2, label2)
        # Only add page break if there are more rooms after this pair
        if idx + 2 < len(room_pngs):
            doc.add_page_break()

    out_path = Path(args.output).resolve() if args.output else OUTPUT_DOCX
    # Ensure parent directory exists
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")

if __name__ == "__main__":
    main()


